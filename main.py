import os
import logging
from flask import Flask, render_template, request, jsonify, send_file, session
from flask_wtf.csrf import CSRFProtect, generate_csrf
from models import db, Doctor, Template, Report
from dotenv import load_dotenv
from datetime import datetime
from sqlalchemy import create_engine
from sqlalchemy.exc import OperationalError
import time
import html2text
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from bs4 import BeautifulSoup

# Carregar variáveis de ambiente
load_dotenv()

# Configurar logging detalhado
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def get_database_url():
    """Get database URL with proper SSL configuration"""
    database_url = os.environ.get('DATABASE_URL')
    if not database_url:
        raise ValueError("DATABASE_URL é obrigatória")

    # Adicionar parâmetros SSL se não existirem
    if 'sslmode=' not in database_url:
        database_url += '?sslmode=require'
    return database_url

def create_app():
    app = Flask(__name__)

    try:
        # Configuração detalhada do banco de dados
        database_url = get_database_url()
        logger.info("Inicializando aplicação com configurações...")
        logger.info(f"Database URL: {database_url}")

        # Configuração básica
        app.config.update(
            SQLALCHEMY_DATABASE_URI=database_url,
            SQLALCHEMY_TRACK_MODIFICATIONS=False,
            SQLALCHEMY_ENGINE_OPTIONS={
                'pool_pre_ping': True,
                'pool_recycle': 300,
            },
            SECRET_KEY=os.environ.get('SECRET_KEY', os.urandom(32)),
            WTF_CSRF_ENABLED=True,
            WTF_CSRF_SECRET_KEY=os.environ.get('WTF_CSRF_SECRET_KEY', os.urandom(32)),
            WTF_CSRF_CHECK_DEFAULT=True,
            SESSION_COOKIE_SECURE=True,
            SESSION_COOKIE_HTTPONLY=True,
            REMEMBER_COOKIE_SECURE=True,
            REMEMBER_COOKIE_HTTPONLY=True
        )

        # Inicializar extensões com tratamento de erro detalhado
        logger.info("Inicializando extensões do Flask...")
        db.init_app(app)
        csrf = CSRFProtect()
        csrf.init_app(app)
        logger.info("Extensões inicializadas com sucesso")

        return app
    except Exception as e:
        logger.error(f"Erro crítico ao inicializar aplicação: {str(e)}", exc_info=True)
        raise

app = create_app()

def retry_database_operation(operation, max_retries=3):
    """Retry database operations with exponential backoff"""
    for attempt in range(max_retries):
        try:
            return operation()
        except OperationalError as e:
            if attempt == max_retries - 1:
                raise
            wait_time = 2 ** attempt
            logger.warning(f"Tentativa {attempt + 1} falhou, tentando novamente em {wait_time}s: {str(e)}")
            time.sleep(wait_time)

@app.route('/api/doctors', methods=['POST'])
def create_doctor():
    """Create a new doctor with enhanced error handling and validation"""
    try:
        if not request.is_json:
            logger.error("Request Content-Type is not application/json")
            return jsonify({
                "error": "Content-Type deve ser application/json"
            }), 400

        data = request.get_json()
        logger.info(f"Tentativa de criar médico com dados: {data}")

        # Validação dos campos obrigatórios
        if not data.get('full_name'):
            logger.error("Tentativa de criar médico sem nome")
            return jsonify({"error": "Nome completo é obrigatório"}), 400
        if not data.get('crm'):
            logger.error("Tentativa de criar médico sem CRM")
            return jsonify({"error": "CRM é obrigatório"}), 400

        def check_existing_doctor():
            return Doctor.query.filter_by(crm=data['crm']).first()

        # Validação do CRM existente com retry
        existing_doctor = retry_database_operation(check_existing_doctor)
        if existing_doctor:
            logger.error(f"Tentativa de criar médico com CRM já existente: {data['crm']}")
            return jsonify({"error": "CRM já cadastrado"}), 400

        # Criar novo médico
        new_doctor = Doctor(
            full_name=data['full_name'],
            crm=data['crm'],
            rqe=data.get('rqe', '')
        )

        def save_doctor():
            db.session.add(new_doctor)
            db.session.commit()
            return new_doctor

        # Salvar com retry
        new_doctor = retry_database_operation(save_doctor)
        logger.info(f"Médico criado com sucesso: ID {new_doctor.id}")

        return jsonify({
            "message": "Médico cadastrado com sucesso",
            "doctor": new_doctor.to_dict()
        }), 201

    except Exception as e:
        db.session.rollback()
        logger.error(f"Erro ao criar médico: {str(e)}", exc_info=True)
        return jsonify({"error": f"Erro ao criar médico: {str(e)}"}), 500

@app.route('/')
def index():
    try:
        logger.info("Acessando rota principal")
        doctors = Doctor.query.all()
        templates = Template.query.all()
        logger.info(f"Encontrados {len(doctors)} médicos e {len(templates)} templates")
        return render_template('index.html', doctors=doctors, templates=templates)
    except Exception as e:
        logger.error(f"Erro ao carregar página inicial: {str(e)}", exc_info=True)
        return "Erro ao conectar ao banco de dados", 500

@app.route('/templates')
def templates():
    doctors = Doctor.query.all()
    templates = Template.query.all()
    return render_template('templates.html', doctors=doctors, templates=templates)

@app.route('/doctors')
def doctors():
    try:
        logger.info("Acessando página de médicos")
        doctors = Doctor.query.all()
        logger.info(f"Encontrados {len(doctors)} médicos")
        return render_template('doctors.html', doctors=doctors)
    except Exception as e:
        logger.error(f"Erro ao carregar página de médicos: {str(e)}", exc_info=True)
        return "Erro ao carregar página de médicos", 500

@app.route('/reports')
def reports():
    try:
        logger.info("Acessando página de relatórios")
        doctors = Doctor.query.all()
        templates = Template.query.all()
        logger.info(f"Encontrados {len(doctors)} médicos e {len(templates)} templates")
        return render_template('reports.html', doctors=doctors, templates=templates)
    except Exception as e:
        logger.error(f"Erro ao carregar página de relatórios: {str(e)}", exc_info=True)
        return "Erro ao carregar página de relatórios", 500

@app.route('/api/templates', methods=['GET'])
def get_templates():
    try:
        category = request.args.get('category')
        query = Template.query
        if category:
            query = query.filter_by(category=category)
        templates = query.all()
        return jsonify([template.to_dict() for template in templates])
    except Exception as e:
        logger.error(f"Erro ao buscar templates: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route('/api/templates', methods=['POST'])
def create_template():
    """Create a new template with enhanced error handling and logging"""
    try:
        if not request.is_json:
            logger.error("Requisição não contém Content-Type: application/json")
            return jsonify({"error": "Content-Type deve ser application/json"}), 400

        data = request.get_json()
        logger.debug(f"Dados recebidos para criar template: {data}")

        # Validar dados obrigatórios
        if not data.get('name'):
            logger.error("Tentativa de criar template sem título")
            return jsonify({"error": "Título é obrigatório"}), 400

        if not data.get('content'):
            logger.error("Tentativa de criar template sem conteúdo")
            return jsonify({"error": "Conteúdo é obrigatório"}), 400

        category = data.get('category', 'mascara')
        logger.info(f"Criando novo template com categoria: {category}")

        new_template = Template(
            name=data['name'],
            content=data['content'],
            category=category,
            doctor_id=data.get('doctor_id')
        )

        db.session.add(new_template)
        db.session.commit()
        logger.info(f"Template criado com sucesso: ID {new_template.id}")

        return jsonify({
            "message": "Template salvo com sucesso",
            "template": new_template.to_dict()
        }), 201

    except Exception as e:
        db.session.rollback()
        logger.error(f"Erro ao criar template: {str(e)}", exc_info=True)
        return jsonify({"error": f"Erro ao salvar template: {str(e)}"}), 400

@app.route('/api/doctors', methods=['GET'])
def get_doctors():
    try:
        doctors = Doctor.query.all()
        return jsonify([doctor.to_dict() for doctor in doctors])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/doctors/<int:id>', methods=['PUT'])
def update_doctor(id):
    try:
        doctor = Doctor.query.get_or_404(id)
        data = request.get_json()

        # Validar dados
        if not data.get('full_name'):
            return jsonify({"error": "Nome completo é obrigatório"}), 400
        if not data.get('crm'):
            return jsonify({"error": "CRM é obrigatório"}), 400

        # Verificar se o CRM já existe para outro médico
        existing_doctor = Doctor.query.filter(Doctor.crm == data['crm'], Doctor.id != id).first()
        if existing_doctor:
            return jsonify({"error": "CRM já cadastrado para outro médico"}), 400

        # Atualizar dados
        doctor.full_name = data['full_name']
        doctor.crm = data['crm']
        doctor.rqe = data.get('rqe', '')

        db.session.commit()
        return jsonify(doctor.to_dict())
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400

@app.route('/api/doctors/<int:id>', methods=['DELETE'])
def delete_doctor(id):
    try:
        doctor = Doctor.query.get_or_404(id)

        # Verificar se o médico tem laudos ou templates associados
        if doctor.reports or doctor.templates:
            return jsonify({
                "error": "Não é possível excluir o médico pois existem laudos ou templates associados"
            }), 400

        db.session.delete(doctor)
        db.session.commit()
        return jsonify({"message": "Médico excluído com sucesso"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400


@app.route('/api/reports', methods=['POST'])
def create_report():
    try:
        data = request.get_json()

        # Converter string de data para objeto date
        exam_date = datetime.strptime(data['exam_date'], '%Y-%m-%d').date() if data.get('exam_date') else None
        birthdate = datetime.strptime(data['patient_birthdate'], '%Y-%m-%d').date() if data.get('patient_birthdate') else None

        new_report = Report(
            patient_name=data['patient_name'],
            patient_birthdate=birthdate,
            patient_gender=data['patient_gender'],
            patient_weight=data.get('patient_weight'),
            patient_height=data.get('patient_height'),
            body_surface=data.get('body_surface'),
            exam_date=exam_date,
            left_atrium=data.get('left_atrium'),
            aorta=data.get('aorta'),
            ascending_aorta=data.get('ascending_aorta'),
            diastolic_diameter=data.get('diastolic_diameter'),
            systolic_diameter=data.get('systolic_diameter'),
            septum_thickness=data.get('septum_thickness'),
            posterior_wall=data.get('posterior_wall'),
            right_ventricle=data.get('right_ventricle'),
            content=data['content'],
            doctor_id=data['doctor_id']
        )

        db.session.add(new_report)
        db.session.commit()
        return jsonify(new_report.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400

@app.route('/gerar_doc', methods=['POST'])
def gerar_doc():
    try:
        data = request.get_json()
        logger.info("Iniciando geração do documento DOC")
        logger.debug("Dados recebidos para geração do DOC: %s", data)

        doc = Document()
        logger.debug("Documento DOC iniciado")

        # Configuração inicial do documento - Ajuste de margens e espaçamento
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Título com espaçamento reduzido
        title = doc.add_heading('Laudo de Ecodopplercardiograma', 0)
        title.paragraph_format.space_after = Pt(12)

        # Dados do Paciente com espaçamento otimizado
        patient_heading = doc.add_heading('Dados do Paciente', level=1)
        patient_heading.paragraph_format.space_before = Pt(6)
        patient_heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        logger.debug("Tabela inicial criada")

        # Cabeçalhos da tabela
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Campo'
        header_cells[1].text = 'Valor'

        # Dados do paciente
        paciente = data.get('paciente', {})
        logger.info("Processando dados do paciente: %s", paciente)

        dados_paciente = [
            ('Nome', paciente.get('nome', 'N/D')),
            ('Data de Nascimento', paciente.get('dataNascimento', 'N/D')),
            ('Sexo', paciente.get('sexo', 'N/D')),
            ('Peso', paciente.get('peso', 'N/D')),
            ('Altura', paciente.get('altura', 'N/D')),
            ('Data do Exame', paciente.get('dataExame', 'N/D'))
        ]

        for campo, valor in dados_paciente:
            row_cells = table.add_row().cells
            row_cells[0].text = campo
            row_cells[1].text = str(valor)

        logger.debug("Dados do paciente adicionados à tabela")

        # Medidas e Cálculos com espaçamento otimizado
        measures_heading = doc.add_heading('Medidas e Cálculos', level=1)
        measures_heading.paragraph_format.space_before = Pt(12)
        measures_heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = True

        # Cabeçalhos
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Medida'
        header_cells[1].text = 'Valor'
        header_cells[2].text = 'Cálculo'
        header_cells[3].text = 'Resultado'

        medidas = data.get('medidas', {})
        calculos = data.get('calculos', {})
        logger.info("Processando medidas: %s", medidas)
        logger.info("Processando cálculos: %s", calculos)

        medidas_calculos = [
            ('Átrio Esquerdo', medidas.get('atrio', 'N/D'), 
             'Volume Diastólico Final', calculos.get('volumeDiastFinal', 'N/D')),
            ('Aorta', medidas.get('aorta', 'N/D'), 
             'Volume Sistólico Final', calculos.get('volumeSistFinal', 'N/D')),
            ('Diâmetro Diastólico', medidas.get('diamDiastFinal', 'N/D'), 
             'Volume Ejetado', calculos.get('volumeEjetado', 'N/D')),
            ('Diâmetro Sistólico', medidas.get('diamSistFinal', 'N/D'), 
             'Fração de Ejeção', calculos.get('fracaoEjecao', 'N/D')),
            ('Espessura do Septo', medidas.get('espDiastSepto', 'N/D'),
             'Percentual Enc. Cavidade', calculos.get('percentEncurt', 'N/D')),
            ('Espessura PPVE', medidas.get('espDiastPPVE', 'N/D'),
             'Espessura Relativa', calculos.get('espRelativa', 'N/D')),
            ('Ventrículo Direito', medidas.get('vd', 'N/D'),
             'Massa do VE', calculos.get('massaVE', 'N/D'))
        ]

        for medida, valor_medida, calculo, valor_calculo in medidas_calculos:
            row_cells = table.add_row().cells
            row_cells[0].text = medida
            row_cells[1].text = str(valor_medida)
            row_cells[2].text = calculo
            row_cells[3].text = str(valor_calculo)

        logger.debug("Medidas e cálculos adicionados à tabela")

        # Laudo com espaçamento otimizado
        laudo_heading = doc.add_heading('Laudo', level=1)
        laudo_heading.paragraph_format.space_before = Pt(12)
        laudo_heading.paragraph_format.space_after = Pt(6)

        laudo_texto = data.get('laudo', '')
        if laudo_texto:
            # Configurar HTML2Text para preservar apenas quebras de linha necessárias
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.body_width = 0  # Desativar quebra automática de linha
            h.ignore_emphasis = True
            h.ignore_tables = True
            h.single_line_break = True  # Usar apenas uma quebra de linha

            # Converter HTML para texto
            laudo_texto = h.handle(laudo_texto)

            # Limpar espaços extras e quebras de linha desnecessárias
            laudo_texto = '\n'.join(line.strip() for line in laudo_texto.splitlines() if line.strip())

            # Adicionar parágrafos com espaçamento reduzido
            for paragrafo in laudo_texto.split('\n'):
                if paragrafo.strip():
                    p = doc.add_paragraph(paragrafo.strip())
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.0

            logger.debug("Laudo adicionado ao documento")
        else:
            logger.warning("Laudo vazio recebido")

        # Assinatura do Médico com espaçamento otimizado
        if data.get('medico'):
            logger.info("Processando dados do médico: %s", data['medico'])
            medico = data['medico']

            # Validação mais robusta dos dados do médico
            nome_medico = medico.get('nome', '').strip()
            crm = medico.get('crm', '').strip()
            rqe = medico.get('rqe', '').strip()

            logger.debug(f"Dados do médico extraídos - Nome: {nome_medico}, CRM: {crm}, RQE: {rqe}")

            if nome_medico and crm:
                doc.add_paragraph()  # Espaço antes da assinatura

                # Linha para assinatura
                linha = doc.add_paragraph('_' * 40)
                linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
                linha.paragraph_format.space_after = Pt(6)

                # Assinatura com formatação melhorada
                assinatura_para = doc.add_paragraph()
                assinatura_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                assinatura_para.paragraph_format.space_before = Pt(6)
                assinatura_para.paragraph_format.space_after = Pt(0)

                # Nome do médico
                run = assinatura_para.add_run(f"Dr. {nome_medico}")
                run.bold = True
                assinatura_para.add_run('\n')

                # CRM e RQE
                crm_text = f"CRM: {crm}"
                if rqe:
                    crm_text += f" / RQE: {rqe}"
                assinatura_para.add_run(crm_text)

                logger.debug("Assinatura do médico adicionada com sucesso")
            else:
                logger.warning(f"Dados incompletos do médico - Nome: {nome_medico}, CRM: {crm}")
                return jsonify({"error": "Dados do médico incompletos"}), 400

        # Salvar documento
        logger.debug("Preparando para salvar o documento")
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        logger.info("Documento DOC gerado com sucesso")

        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='laudo_ecocardiograma.docx'
        )

    except Exception as e:
        logger.error("Erro ao gerar documento DOC: %s", str(e), exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route('/gerar_pdf', methods=['POST'])
def gerar_pdf():
    try:
        data = request.get_json()
        logger.info("Iniciando geração do PDF")
        logger.debug("Dados recebidos: %s", data)

        # Criar PDF em memória
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30
        )
        story.append(Paragraph('Laudo de Ecodopplercardiograma', title_style))

        # Dados do Paciente
        paciente = data.get('paciente', {})
        logger.info("Processando dados do paciente: %s", paciente)
        dados_paciente = [
            ['Campo', 'Valor'],
            ['Nome', paciente.get('nome', 'N/D')],
            ['Data de Nascimento', paciente.get('dataNascimento', 'N/D')],
            ['Sexo', paciente.get('sexo', 'N/D')],
            ['Peso', paciente.get('peso', 'N/D')],
            ['Altura', paciente.get('altura', 'N/D')],
            ['Data do Exame', paciente.get('dataExame', 'N/D')]
        ]

        t = Table(dados_paciente)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))

        # Medidas e Cálculos
        story.append(Paragraph('Medidas e Cálculos', styles['Heading2']))
        medidas = data.get('medidas', {})
        calculos = data.get('calculos', {})
        logger.info("Processando medidas e cálculos")

        medidas_calculos = [
            ['Medida', 'Valor', 'Cálculo', 'Resultado'],
            ['Átrio Esquerdo', medidas.get('atrio', 'N/D'), 
             'Volume Diastólico Final', calculos.get('volumeDiastFinal', 'N/D')],
            ['Aorta', medidas.get('aorta', 'N/D'), 
             'Volume Sistólico Final', calculos.get('volumeSistFinal', 'N/D')],
            ['Diâmetro Diastólico', medidas.get('diamDiastFinal', 'N/D'), 
             'Volume Ejetado', calculos.get('volumeEjetado', 'N/D')],
            ['Diâmetro Sistólico', medidas.get('diamSistFinal', 'N/D'), 
             'Fração de Ejeção', calculos.get('fracaoEjecao', 'N/D')],
            ['Espessura do Septo', medidas.get('espDiastSepto', 'N/D'),
             'Percentual Enc. Cavidade', calculos.get('percentEncurt', 'N/D')],
            ['Espessura PPVE', medidas.get('espDiastPPVE', 'N/D'),
             'Espessura Relativa', calculos.get('espRelativa', 'N/D')],
            ['Ventrículo Direito', medidas.get('vd', 'N/D'),
             'Massa do VE', calculos.get('massaVE', 'N/D')]
        ]

        t = Table(medidas_calculos)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))

        # Laudo
        story.append(Paragraph('Laudo', styles['Heading2']))
        h = html2text.HTML2Text()
        h.ignore_links = True
        laudo_texto = h.handle(data.get('laudo', ''))
        story.append(Paragraph(laudo_texto, styles['Normal']))

        # Assinatura do Médico
        if data.get('medico'):
            story.append(Spacer(1, 30))
            medico = data['medico']
            if medico.get('nome') and medico.get('crm'):
                assinatura = f"{medico['nome']}\nCRM: {medico['crm']}"
                if medico.get('rqe'):
                    assinatura += f"\nRQE: {medico['rqe']}"
                story.append(Paragraph(assinatura, styles['Normal']))

        # Gerar PDF
        doc.build(story)
        buffer.seek(0)
        logger.info("PDF gerado com sucesso")

        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='laudo_ecocardiograma.pdf'
        )

    except Exception as e:
        logger.error(f"Erro ao gerar PDF: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@app.route('/api/phrases', methods=['POST'])
def save_phrase():
    """Save a new phrase or template"""
    try:
        if not request.is_json:
            logger.error("Requisição não contém Content-Type: application/json")
            return jsonify({"error": "Content-Type deve ser application/json"}), 400

        data = request.get_json()
        logger.debug(f"Dados recebidos para salvar frase: {data}")

        if not data.get('content'):
            logger.error("Tentativa de salvar frase sem conteúdo")
            return jsonify({"error": "Conteúdo é obrigatório"}), 400

        if not data.get('title'):
            logger.error("Tentativa de salvar frase sem título")
            return jsonify({"error": "Título é obrigatório"}), 400

        new_template = Template(
            name=data['title'],
            content=data['content'],
            category='mascara',  # Sempre salvar como máscara
            doctor_id=data.get('doctor_id')
        )

        db.session.add(new_template)
        db.session.commit()
        logger.info(f"Nova máscara salva com sucesso: ID {new_template.id}")

        return jsonify({
            "message": "Máscara salva com sucesso",
            "template": new_template.to_dict()
        }), 201

    except Exception as e:
        db.session.rollback()
        logger.error(f"Erro ao salvar máscara: {str(e)}", exc_info=True)
        return jsonify({"error": str(e)}), 400

@app.route('/api/phrases', methods=['GET'])
def list_phrases():
    try:
        phrases = Template.query.filter_by(category='frase').all()
        return jsonify([phrase.to_dict() for phrase in phrases])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/phrases/<int:id>', methods=['DELETE'])
def delete_phrase(id):
    try:
        phrase = Template.query.filter_by(id=id, category='frase').first_or_404()
        db.session.delete(phrase)
        db.session.commit()
        return jsonify({"message": "Frase excluída com sucesso"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400


@app.route('/api/templates/<int:id>', methods=['DELETE'])
def delete_template(id):
    """Delete a template with enhanced error handling and logging"""
    try:
        template = Template.query.get_or_404(id)
        template_name = template.name  # Store name for logging

        db.session.delete(template)
        db.session.commit()

        logger.info(f"Template '{template_name}' (ID: {id}) excluído com sucesso")
        return jsonify({"message": "Template excluído com sucesso"}), 200

    except Exception as e:
        db.session.rollback()
        logger.error(f"Erro ao excluir template {id}: {str(e)}", exc_info=True)
        return jsonify({"error": f"Erro ao excluir template: {str(e)}"}), 400

@app.route('/phrases')
def phrases():
    return render_template('phrases.html')

@app.route('/api/templates/<int:id>/pdf', methods=['GET'])
def export_template_pdf(id):
    """Export a template as PDF with enhanced error handling and logging"""
    try:
        template = Template.query.get_or_404(id)
        logger.info(f"Iniciando exportação do template {id} para PDF")
        logger.debug(f"Dados dotemplate: nome={template.name}, categoria={template.category}")

        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30
        )
        story.append(Paragraph(template.name, title_style))
        story.append(Spacer(1, 20))
        logger.debug("Título do PDF adicionado com sucesso")


        # Content
        # Pre-process HTML with BeautifulSoup
        soup = BeautifulSoup(template.content, 'html.parser')

        # Clean up HTML
        for tag in soup.find_all(['script', 'style']):
            tag.decompose()

        # Convert to markdown
        h = html2text.HTML2Text()
        h.ignore_links = True
        h.unicode_snob = True
        h.body_width = 0
        content_text = h.handle(str(soup))

        # Ensure content is properly sanitized
        if not content_text.strip():
            raise ValueError("Conteúdo do template está vazio após conversão")

        # Split into paragraphs and add to story
        paragraphs = content_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 12))

        logger.debug("Conteúdo do PDF processado com sucesso")

        # Add metadata
        story.append(Spacer(1, 30))
        metadata = [
            f"Categoria: {template.category}",
            f"Criado em: {template.created_at.strftime('%d/%m/%Y %H:%M')}"
        ]
        for meta in metadata:
            story.append(Paragraph(meta, styles['Normal']))
        logger.debug("Metadados do PDF adicionados do PDF adicionados com sucesso")

        # Generate PDF
        doc.build(story)
        buffer.seek(0)
        logger.info(f"PDF do template {id} gerado com sucesso")

        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{template.name}.pdf"
        )

    except Exception as e:
        error_msg = f"Erro ao exportar template {id} para PDF: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return jsonify({"error": error_msg}), 500

@app.route('/api/templates/<int:id>/doc', methods=['GET'])
def export_template_doc(id):
    """Export a template as DOCX with enhanced error handling and logging"""
    try:
        template = Template.query.get_or_404(id)
        logger.info(f"Iniciando exportação do template {id} para DOCX")
        logger.debug(f"Dados do template: nome={template.name}, categoria={template.category}")

        # Create DOCX in memory
        doc = Document()
        doc.add_heading(template.name, 0)
        logger.debug("Documento DOCX criado com título")

        # Add content
        # Pre-process HTML with BeautifulSoup
        soup = BeautifulSoup(template.content, 'html.parser')

        # Clean up HTML
        for tag in soup.find_all(['script', 'style']):
            tag.decompose()

        # Convert to markdown
        h = html2text.HTML2Text()
        h.ignore_links = True
        h.unicode_snob = True
        h.body_width = 0
        content_text = h.handle(str(soup))

        # Ensure content is properly sanitized
        if not content_text.strip():
            raise ValueError("Conteúdo do template está vazio após conversão")

        # Split into paragraphs and add to document
        paragraphs = content_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        logger.debug("Conteúdo do DOCX processado com sucesso")

        # Add metadata
        doc.add_paragraph()  # Add space
        doc.add_paragraph(f"Categoria: {template.category}")
        doc.add_paragraph(f"Criado em: {template.created_at.strftime('%d/%m/%Y %H:%M')}")
        logger.debug("Metadados do DOCX adicionados com sucesso")

        # Save to buffer
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        logger.info(f"DOCX do template {id} gerado com sucesso")

        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{template.name}.docx"
        )

    except Exception as e:
        error_msg = f"Erro ao exportar template {id} para DOCX: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return jsonify({"error": error_msg}), 500

@app.after_request
def add_csrf_header(response):
    """Add CSRF token to response headers and cookies"""
    if 'text/html' in response.headers.get('Content-Type', ''):
        csrf_token = generate_csrf()
        response.headers.set('X-CSRF-Token', csrf_token)
        response.set_cookie('csrf_token', csrf_token, secure=True, httponly=True, samesite='Strict')
    return response

if __name__ == '__main__':
    try:
        logger.info("Tentando criar tabelas do banco de dados...")
        with app.app_context():
            db.create_all()
        logger.info("Tabelas do banco de dados criadas com sucesso")
        app.run(host='0.0.0.0', port=3001, debug=True)
    except Exception as e:
        logger.error(f"Erro ao inicializar aplicação: {str(e)}", exc_info=True)
        raise